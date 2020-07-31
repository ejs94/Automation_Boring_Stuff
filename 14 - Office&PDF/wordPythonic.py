import os
import docx

# Example 1
# Writing in Docx File

os.chdir('.\\14 - Office&PDF')
d = docx.Document('demo.docx')
# d.paragraphs[3].text
d.add_paragraph('Jujubas')
d.save('demo2.docx')

# Example 2
# Read a Docx File


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for p in doc.paragraphs:
        fullText.append(p.text)
    return '\n'.join(fullText)


print(getText('demo.docx'))
